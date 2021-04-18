import os
import xlrd
import pprint
import xlsxwriter
import pandas as pd
import matplotlib.pyplot as plt

from docx.shared import Pt
from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'

book = xlrd.open_workbook('Accenture_Датасет.xlsx')
df1 = pd.read_excel(os.path.join('Accenture_Датасет.xlsx'), engine='openpyxl')
things = {}

for thing, kg, buy, sell in zip(df1["Код товара"][:2470], df1["Продажи в кг"][:2470], df1["Сумма в ценах закупки"][:2470], df1["Сумма в ценах продажи"][:2470]):
    if thing not in things:
        things[thing] = [1, kg, buy, sell]
    else:
        things[thing][0] += 1       # количество
        things[thing][1] += kg      # кг
        things[thing][2] += buy     # цена покупки
        things[thing][3] += sell    # цена продажи

agents_f = {}

for agent, price, buy, kg in zip(df1["Агент"][:2470], df1["Сумма в ценах закупки"][:2470], df1["Сумма в ценах продажи"][:2470], df1["Продажи в кг"][:2470]):
    if agent not in agents_f:
        agents_f[agent] = [price, buy, kg]
    else:
        agents_f[agent][0] += price  # ценв покупки
        agents_f[agent][1] += buy    # цена продажи
        agents_f[agent][2] += kg     # кг

doc.add_heading("Отчёт за 2015 (1 и 2 месяц)", 0)
for i in things.items():
    doc.add_heading("Товар с маркеровкой\t" + str(i[0]), 1)
    doc.add_paragraph(f"Количество проданных единиц \t\t{i[1][0]}", style='List Bullet')
    doc.add_paragraph(f"Количество проданных киллограм \t\t{i[1][1]}", style='List Bullet')
    doc.add_paragraph(f"Общая сумма закупки \t\t\t\t{int(i[1][2])}", style='List Bullet')
    doc.add_paragraph(f"Общая сумма продаж \t\t\t\t{int(i[1][3])}", style='List Bullet')
    doc.add_paragraph(f"Чистая прибыль товара \t\t\t\t{int(i[1][3] - i[1][2])}", style='List Bullet')
doc.add_heading("Всего\t", 1)
doc.add_paragraph(f"Количество проданных единиц \t\t{int(list(things.items())[0][1][0] + list(things.items())[1][1][0] + list(things.items())[2][1][0])}", style='List Bullet')
doc.add_paragraph(f"Количество проданных киллограм \t\t{int(list(things.items())[0][1][1] + list(things.items())[1][1][1] + list(things.items())[2][1][1])}", style='List Bullet')
ras = int(list(things.items())[0][1][2] + list(things.items())[1][1][2] + list(things.items())[2][1][2])
sell = int(list(things.items())[0][1][3] + list(things.items())[1][1][3] + list(things.items())[2][1][3])
doc.add_paragraph(f"Общая сумма закупки \t\t\t\t{ras}", style='List Bullet')
doc.add_paragraph(f"Общая сумма продаж \t\t\t\t{sell}", style='List Bullet')
doc.add_paragraph(f"Чистая прибыль \t\t\t\t\t{sell - ras}", style='List Bullet')
a = '\n\n\n\n\n\n\n'
paragraph = doc.add_paragraph(a)
paragraph.style = doc.styles['Normal']
things_2 = {}

for thing, kg, buy, sell in zip(df1["Код товара"][3857:], df1["Продажи в кг"][3857:], df1["Сумма в ценах закупки"][3857:], df1["Сумма в ценах продажи"][3857:]):
    if thing not in things_2:
        things_2[thing] = [1, kg, buy, sell]
    else:
        things_2[thing][0] += 1       # количество
        things_2[thing][1] += kg      # кг
        things_2[thing][2] += buy     # цена покупки
        things_2[thing][3] += sell    # цена продажи

doc.add_heading("Отчёт за 2016 (1 и 2 месяц)", 0)
sr_key = sorted(things, key=things.get)
sort_1, sort_2 = things, things_2
things, things_2 = {}, {}
for w in sr_key:
    things[w] = sort_1[w]
    things_2[w] = sort_2[w]

for i in things_2.items():
    doc.add_heading("Товар с маркеровкой\t" + str(i[0]), 1)
    doc.add_paragraph(f"Количество проданных единиц \t\t{i[1][0]}", style='List Bullet')
    doc.add_paragraph(f"Количество проданных киллограм \t\t{i[1][1]}", style='List Bullet')
    doc.add_paragraph(f"Общая сумма закупки \t\t\t\t{int(i[1][2])}", style='List Bullet')
    doc.add_paragraph(f"Общая сумма продаж \t\t\t\t{int(i[1][3])}", style='List Bullet')
    doc.add_paragraph(f"Чистая прибыль товара \t\t\t\t{int(i[1][3] - i[1][2])}", style='List Bullet')
doc.add_heading("Всего\t", 1)
doc.add_paragraph(f"Количество проданных единиц \t\t{int(list(things_2.items())[0][1][0] + list(things_2.items())[1][1][0] + list(things_2.items())[2][1][0])}", style='List Bullet')
doc.add_paragraph(f"Количество проданных киллограм \t\t{int(list(things_2.items())[0][1][1] + list(things_2.items())[1][1][1] + list(things_2.items())[2][1][1])}", style='List Bullet')
ras = int(list(things_2.items())[0][1][2] + list(things_2.items())[1][1][2] + list(things_2.items())[2][1][2])
sell = int(list(things_2.items())[0][1][3] + list(things_2.items())[1][1][3] + list(things_2.items())[2][1][3])
doc.add_paragraph(f"Общая сумма закупки \t\t\t\t{ras}", style='List Bullet')
doc.add_paragraph(f"Общая сумма продаж \t\t\t\t{sell}", style='List Bullet')
doc.add_paragraph(f"Чистая прибыль \t\t\t\t\t{sell - ras}", style='List Bullet')
a = '\n\n\n\n\n\n\n'
paragraph = doc.add_paragraph(a)
paragraph.style = doc.styles['Normal']

agents_s = {}

for agent, price, buy, kg in zip(df1["Агент"][3857:], df1["Сумма в ценах закупки"][3857:], df1["Сумма в ценах продажи"][3857:], df1["Продажи в кг"][3857:]):
    if agent not in agents_s:
        agents_s[agent] = [price, buy, kg]
    else:
        agents_s[agent][0] += price  # ценв покупки
        agents_s[agent][1] += buy    # цена продажи
        agents_s[agent][2] += kg     # kg

all_buy = 0
for price in df1["Сумма в ценах закупки"][:3857]:
    all_buy += float(price)
all_sell = 0
for price in df1["Сумма в ценах продажи"][:3857]:
    all_sell += float(price)

doc.add_heading("Отчёт об агентах за 2015 (1 и 2 месяц)", 0)
for i in range(len(list(agents_f.items()))):
    dohod = int((list(agents_f.items())[i][1][1] - list(agents_f.items())[i][1][0]))
    a = "2015 - 2016 Чистая прибыль за кг от   " + list(agents_f.items())[i][0].upper().ljust(20, ' ') + str(int((list(agents_f.items())[i][1][1] - list(agents_f.items())[i][1][0]) / list(agents_f.items())[i][1][2])).rjust(5, ' ') + '%'
    b = "\t\tЗакупка агента:\t\t" + str(int(list(agents_f.items())[i][1][0])).rjust(8, ' ') + "\tот общего V\t" + str(round(list(agents_f.items())[i][1][0] / all_buy * 100, 2)) + '%'
    c = "\t\tПрибыль агента:\t\t" + str(int(list(agents_f.items())[i][1][1])).rjust(8, ' ') + "\tот общего V\t" + str(round(list(agents_f.items())[i][1][1] / all_sell * 100, 2)) + '%'
    d = "\t\tЧистый доход:\t\t" + str(dohod).rjust(8, ' ') + "\tот общего V\t" + str(round(dohod / (all_sell - all_buy) * 100, 2)) + '%'
    doc.add_heading(f"{a}", 1)
    doc.add_paragraph(f"{b}", style='Normal')
    doc.add_paragraph(f"{c}", style='Normal')
    doc.add_paragraph(f"{d}", style='Normal')
a = '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n'
paragraph = doc.add_paragraph(a)
paragraph.style = doc.styles['Normal']

doc.add_heading("Отчёт об агентах за 2016 (1 и 2 месяц)", 0)
for i in range(len(list(agents_s.items()))):
    dohod = int((list(agents_s.items())[i][1][1] - list(agents_s.items())[i][1][0]))
    a = "2015 - 2016 Чистая прибыль за кг от   " + list(agents_s.items())[i][0].upper().ljust(20, ' ') + str(int(
        (list(agents_s.items())[i][1][1] - list(agents_s.items())[i][1][0]) / list(agents_s.items())[i][1][2])).rjust(5,' ') + '%'
    b = "\t\tЗакупка агента:\t\t" + str(int(list(agents_s.items())[i][1][0])).rjust(8, ' ') + "\tот общего V\t" + str(
        round(list(agents_s.items())[i][1][0] / all_buy * 100, 2)) + '%'
    c = "\t\tПрибыль агента:\t\t" + str(int(list(agents_s.items())[i][1][1])).rjust(8, ' ') + "\tот общего V\t" + str(
        round(list(agents_s.items())[i][1][1] / all_sell * 100, 2)) + '%'
    d = "\t\tЧистый доход:\t\t" + str(dohod).rjust(8, ' ') + "\tот общего V\t" + str(
        round(dohod / (all_sell - all_buy) * 100, 2)) + '%'
    doc.add_heading(f"{a}", 1)
    doc.add_paragraph(f"{b}", style='Normal')
    doc.add_paragraph(f"{c}", style='Normal')
    doc.add_paragraph(f"{d}", style='Normal')

statik_first = {}
for date, count, money in zip(df1["Дата"][:2470], df1["Продажи в кг"][:2470], df1["Сумма в ценах продажи"][:2470]):
    if date not in statik_first:
        statik_first[date] = [count, money]
    else:
        statik_first[date][0] += count      # количество в кг
        statik_first[date][1] += money      # продажа

statik_second = {}
for date, count, money in zip(df1["Дата"][3858:], df1["Продажи в кг"][3858:], df1["Сумма в ценах продажи"][3858:]):
    if date not in statik_second:
        statik_second[date] = [count, money]
    else:
        statik_second[date][0] += count  # количество в кг
        statik_second[date][1] += money  # продажа

date_1, summa_kg_1, summa_1 = [], [], []
date_2, summa_kg_2, summa_2 = [], [], []

for i in sorted(zip(statik_first.keys(), list(statik_first.values())), key=lambda x: (x[0].split('.')[1], x[0].split('.')[0])):
    date_1.append(i[0])
    summa_kg_1.append(i[1][0])
    summa_1.append(i[1][1])

for i in sorted(zip(statik_second.keys(), list(statik_second.values())), key=lambda x: (x[0].split('.')[1], x[0].split('.')[0])):
    date_2.append(i[0])
    summa_kg_2.append(i[1][0])
    summa_2.append(i[1][1])

doc.add_heading("Изменения 2015 (1 и 2)  2016 (1 и 2)", 0)
a = "Изменения продажи товаров (в кг) составил:    \t\t" + str(int((sum(summa_kg_2) - sum(summa_kg_1)) / sum(summa_kg_1) * 10000) / 100) + "%"
b = "Изменения выручки с продажи товаров составил: \t" + str(int((sum(summa_2)- sum(summa_1)) / sum(summa_1) * 10000) / 100) + "%"
doc.add_paragraph(f"{a}", style='Normal')
doc.add_paragraph(f"{b}", style='Normal')
a = '\n'
paragraph = doc.add_paragraph(a)
paragraph.style = doc.styles['Normal']

for i in range(3):
    f = list(things.values())[i][3] / list(things.values())[i][1]
    s = list(things_2.values())[i][3] / list(things_2.values())[i][1]
    a = "Изменения цен товара " + str(list(things_2.keys())[i]) + " за кг составили: \t" + str(int((s - f) / f * 10000) / 100) + "%"
    doc.add_paragraph(f"{a}", style='Normal')

a = '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n'
paragraph = doc.add_paragraph(a)
paragraph.style = doc.styles['Normal']
a = '\n'
paragraph = doc.add_paragraph(a)
paragraph.style = doc.styles['Normal']
doc.add_heading('Анализ', 0)

third = {}
for thing, kg, buy, sell in zip(df1["Код товара"][2470:3857], df1["Продажи в кг"][2470:3857], df1["Сумма в ценах закупки"][2470:3857], df1["Сумма в ценах продажи"][2470:3857]):
    if thing not in third:
        third[thing] = [1, kg, buy, sell]
    else:
        third[thing][0] += 1       # количество
        third[thing][1] += kg      # кг
        third[thing][2] += buy     # цена покупки
        third[thing][3] += sell    # цена продажи


for i in range(3):
    doc.add_heading('Товар с маркером ' + str(list(things_2.keys())[i]), 1)
    doc.add_heading('Анализ 2015(1 и 2) : 2016(1 и 2) ', 2)
    f = list(things.values())[i][3] - list(things.values())[i][2]
    s = list(things_2.values())[i][3] - list(things_2.values())[i][2]

    th = list(third.values())[i][3] - list(third.values())[i][2]

    a = "Изменения прибыли товара 2015(1 и 2) : 2016(1 и 2) составили: \t" + str(int((s - f) / f * 10000) / 100) + "%"
    doc.add_paragraph(f"{a}", style='Normal')
    f = list(things.values())[i][1]
    s = list(things_2.values())[i][1]
    a = "Изменения спроса на товар 2015(1 и 2) : 2016(1 и 2) составили: \t" + str(int((s - f) / f * 10000) / 100) + "%"
    doc.add_paragraph(f"{a}", style='Normal')
    doc.add_heading('Прогноз на 2016(3) ', 2)
    if f < s:
        f_p = list(things.values())[i][3] / list(things.values())[i][1]
        s_p = list(things_2.values())[i][3] / list(things_2.values())[i][1]
        th_p = list(third.values())[i][3] / list(third.values())[i][1]

        f_kg = list(things.values())[i][1]
        s_kg = list(things_2.values())[i][1]
        th_kg = list(third.values())[i][1]

        f_buy = list(things.values())[i][3]
        s_buy = list(things_2.values())[i][3]
        th_buy = list(third.values())[i][3]

        k_kg = abs(round((s_kg - f_kg) / f_kg * 100, 2))
        k_buy = abs(round((s_buy - f_buy) / f_buy * 100, 2))

        b = "Рекомендуемая цена на 2016(3):  ".ljust(50) + '\t\t' + str(round(th_p + (s_p-f_p) / f_p * 100, 2))
        c = "Прогнозируемый доход 2016(3):   ".ljust(50) + '\t\t' + str(int(th_buy + s_buy / 2 * (k_buy / 2 / 100)))
        d = "Примерный объём закупок 2016(3):".ljust(50) + '\t\t' + str(int(th_kg + s_kg / 2 * (k_kg / 2 / 100)))
        doc.add_paragraph(f"{b}", style='Normal')
        doc.add_paragraph(f"{c}", style='Normal')
        doc.add_paragraph(f"{d}", style='Normal')

    elif f >= s:
        f_p = list(things.values())[i][3] / list(things.values())[i][1]
        s_p = list(things_2.values())[i][3] / list(things_2.values())[i][1]
        th_p = list(third.values())[i][3] / list(third.values())[i][1]

        f_kg = list(things.values())[i][1]
        s_kg = list(things_2.values())[i][1]
        th_kg = list(third.values())[i][1]

        f_buy = list(things.values())[i][3]
        s_buy = list(things_2.values())[i][3]
        th_buy = list(third.values())[i][3]

        k_kg = abs(round((s_kg - f_kg) / f_kg * 100, 2))
        k_buy = abs(round((s_buy - f_buy) / f_buy * 100, 2))

        b = "Рекомендуемая цена на 2016(3):  ".ljust(50) + '\t\t' + str(round(th_p - (s_p - f_p) / f_p * 100, 2))
        c = "Прогнозируемый доход 2016(3):   ".ljust(50) + '\t\t' + str(int(th_buy + s_buy / 2 * (k_buy / 2 / 100)))
        d = "Примерный объём закупок 2016(3):".ljust(50) + '\t\t' + str(int(th_kg + s_kg / 2 * (k_kg / 2 / 100)))
        doc.add_paragraph(f"{b}", style='Normal')
        doc.add_paragraph(f"{c}", style='Normal')
        doc.add_paragraph(f"{d}", style='Normal')

doc.add_heading('Изменения спроса при рекомендуемых параметрах', 0)
for i in range(3):
    f_p = list(things.values())[i][3] / list(things.values())[i][1]
    s_p = list(things_2.values())[i][3] / list(things_2.values())[i][1]
    th_p = list(third.values())[i][3] / list(third.values())[i][1]

    f_kg = list(things.values())[i][1]
    s_kg = list(things_2.values())[i][1]
    th_kg = list(third.values())[i][1]

    f_buy = list(things.values())[i][3]
    s_buy = list(things_2.values())[i][3]
    th_buy = list(third.values())[i][3]

    k_kg = abs(round((s_kg - f_kg) / f_kg * 100, 2))
    k_buy = abs(round((s_buy - f_buy) / f_buy * 100, 2))

    a = "На товар " + str(list(things_2.keys())[i]) + " составят: \t\t" + str(k_kg) + "%   \tна 2017(1, 2)"
    doc.add_paragraph(f"{a}", style='Normal')

df = pd.DataFrame({'Месяц': [1, 2, 3],
                    'Код товара': [list(things.keys())[0], list(things.keys())[1], list(things.keys())[2]],
                    'Продажи в 2015': [list(things.values())[0][0], list(things.values())[1][0], list(things.values())[2][0]],
                    'Продажи в 2016': [list(things_2.values())[0][0], list(things_2.values())[1][0], list(things_2.values())[2][0]]})

# ------------------------------------------------Excel--------------------------------------------------------------

ex = {}
for i in range(3):
    f_p = list(things.values())[i][3] / list(things.values())[i][1]
    s_p = list(things_2.values())[i][3] / list(things_2.values())[i][1]
    th_p = list(third.values())[i][3] / list(third.values())[i][1]

    f_kg = list(things.values())[i][1]
    s_kg = list(things_2.values())[i][1]
    th_kg = list(third.values())[i][1]

    f_buy = list(things.values())[i][3]
    s_buy = list(things_2.values())[i][3]
    th_buy = list(third.values())[i][3]

    k_kg = abs(round((s_kg - f_kg) / f_kg * 100, 2))
    k_buy = abs(round((s_buy - f_buy) / f_buy * 100, 2))

    if f < s:
        ex[str(list(things_2.keys())[i])] = [str(round(th_p + (s_p - f_p) / f_p * 100, 2)), int(th_buy + s_buy / 2 * (k_buy / 2 / 100)),int(th_kg + s_kg / 2 * (k_kg / 2 / 100))]
    else:
        ex[str(list(things_2.keys())[i])] = [str(round(th_p - (s_p - f_p) / f_p * 100, 2)), int(th_buy + s_buy / 2 * (k_buy / 2 / 100)),int(th_kg + s_kg / 2 * (k_kg / 2 / 100))]

workbook = xlsxwriter.Workbook('Агенты.xlsx')
worksheet = workbook.add_worksheet()
for i, j in enumerate([16, 25, 26, 15]):
    worksheet.set_column(i, i, j)
columns = [{'header': 'Агент'}, {'header': "Сумма в ценах закупки"},
           {'header': "Сумма в ценах продажи"}, {'header': "Продажи в кг"}]
data1 = [[a, *i] for a, i in agents_f.items()]
worksheet.write(0, 0, 'Статистика по агентам за первый год(1 и 2 месяц):')
worksheet.add_table(1, 0, len(data1) + 1, 3, {'data': data1, 'columns': columns})
data2 = [[a, *i] for a, i in agents_s.items()]
worksheet.write(len(data1) + 2, 0, 'Статистика по агентам за второй год(1 и 2 месяц):')
worksheet.add_table(len(data1) + 3, 0, len(data2) + len(data1) + 3, 3, {'data': data2, 'columns': columns})
workbook.close()

workbook = xlsxwriter.Workbook('Объемы.xlsx')
worksheet = workbook.add_worksheet()
for i, j in enumerate([13, 24, 25, 29]):
    worksheet.set_column(i, i, j)
columns = [{'header': 'Код товара'}, {'header': 'Количество покупок'}, {'header': "Продажи в кг"},
           {'header': "Сумма в ценах закупки"}, {'header': "Сумма в ценах продажи"}, {'header': "Чистая прибль"}]
sp = {}
for i in list(things.keys()):
    sp[i] = []
    for j in range(4):
        a = things[i][j] + third[i][j]
        sp[i].append(a)
data1 = [[a, *i] for a, i in sp.items()]
for i in range(len(data1)):
    data1[i].append(data1[i][4] - data1[i][3])

worksheet.write(0, 0, 'Обьемы продаж 2015')
worksheet.add_table(1, 0, len(data1) + 1, 5, {'data': data1, 'columns': columns})
data2 = [[a, *i] for a, i in things_2.items()]
for i in range(len(data2)):
    data2[i].append(data2[i][4] - data2[i][3])
worksheet.write(len(data1) + 2, 0, 'Обьемы продаж 2016(1, 2):')
worksheet.add_table(len(data1) + 3, 0, len(data2) + len(data1) + 3, 5, {'data': data2, 'columns': columns})
columns = [{'header': 'Код товара'}, {'header': "Рекомендуемая цена"},
           {'header': "Сумма в ценах продажи"}, {'header': "Примерный объём закупок в кг"}]
data = [[a, *i] for a, i in ex.items()]
worksheet.write(len(data2) + len(data1) + 4, 0, 'Прогноз:')
worksheet.add_table(len(data2) + len(data1) + 5, 0, len(data) + len(data2) + len(data1) + 5, 3, {'data': data, 'columns': columns})
for i in range(3):
    worksheet.write(16 + i, 0, f'=Sheet1!$C${i + 3}')
    worksheet.write(16 + i, 1, f'=Sheet1!$C${i + 8}')
    worksheet.write(16 + i, 2, f'=Sheet1!$D${i + 13}')
chart = workbook.add_chart({'type': 'line'})
chart.add_series({'categories': '={"Первый";"Второй";"Прогноз"}', 'name': f'{list(things.keys())[0]}', 'values': '=Sheet1!$A$17:$C$17'})
chart.add_series({'name': f'{list(things.keys())[1]}', 'values': '=Sheet1!$A$18:$C$18'})
chart.add_series({'name': f'{list(things.keys())[2]}', 'values': '=Sheet1!$A$19:$C$19'})
chart.set_title({'name': 'Статистика по продажам кг'})
worksheet.insert_chart('A17', chart)
workbook.close()
doc.save('Отчет.docx')
